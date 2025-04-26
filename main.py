from flask import Flask, request, send_file, url_for
from docx import Document
import os
import uuid
import google.generativeai as genai
from dotenv import load_dotenv
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# Initialize Flask
app = Flask(__name__)


# Set up limiter
limiter = Limiter(
    get_remote_address,  # identifies the client IP
    app=app,
    default_limits=["10 per hour"],  # global limit (100 requests per hour)

)





# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Ensure 'docs' directory exists
if not os.path.exists('docs'):
    os.makedirs('docs')

# AI summary generation function
def generate_professional_summary(name, role, skills, experience):
    prompt = f"""Write a professional resume summary for:
    Name: {name}
    Role: {role}
    Skills: {', '.join(skills)}
    Experience: {experience}
    
    Format it in 3-4 sentences using a formal and professional tone."""
    
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(prompt)
    return response.text.strip()


@app.route("/limited")
@limiter.limit("10 per minute")  # custom limit on this route
def limited_route():
    return {"message": "You have access to this route."}

@app.route("/open")
def open_route():
    return {"message": "This route has no rate limit."}



@app.route("/")
def default():
    return {"success": True, "message": "Server is Running"}

@app.route("/health")
def health():
    return {"success": True, "message": "Server is healthy"}

@app.route('/docs/<filename>')
def serve_doc(filename):
    return send_file(os.path.join('docs', filename))

@app.route("/resume", methods=["POST"])
def create_resume():
    data = request.json
    doc = Document()

    # Add Name
    doc.add_heading(data.get("name", "Unnamed"), 0)

    # Contact Info
    doc.add_paragraph(f"Email: {data.get('email', '')}")
    doc.add_paragraph(f"Phone: {data.get('phone', '')}")
    doc.add_paragraph(f"Address: {data.get('address', '')}")

    # Professional Summary
    doc.add_heading("Professional Summary", level=1)
    summary = data.get("summary", "")
    if not summary:
        summary = generate_professional_summary(
            name=data.get("name", "Candidate"),
            role=data.get("role", "Professional"),
            skills=data.get("skills", []),
            experience=", ".join([exp.get("description", "") for exp in data.get("experience", [])])
        )
    doc.add_paragraph(summary)

    # Skills
    doc.add_heading("Skills", level=1)
    skills = data.get("skills", [])
    doc.add_paragraph(", ".join(skills))

    # Experience
    doc.add_heading("Experience", level=1)
    for exp in data.get("experience", []):
        doc.add_paragraph(f"{exp['role']} at {exp['company']} ({exp['duration']})", style='List Bullet')
        doc.add_paragraph(exp.get("description", ""), style='Intense Quote')

    # Education
    doc.add_heading("Education", level=1)
    for edu in data.get("education", []):
        doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['year']})", style='List Bullet')

    # Save Resume
    filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
    file_path = os.path.join('docs', filename)
    doc.save(file_path)

    # Return Download Link
    return {
        "success": True,
        "message": "Resume created successfully with AI-generated summary" if not data.get("summary") else "Resume created successfully",
        "download_url": url_for('serve_doc', filename=filename, _external=True)
    }

if __name__ == "__main__":
    app.run(debug=True)
